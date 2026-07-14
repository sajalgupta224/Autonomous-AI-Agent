from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from datetime import datetime
import os
from app.logger import logger


def generate_doc(document_text, user_request=None, tasks=None):
    os.makedirs("generated_docs", exist_ok=True)
    doc = Document()

    # -----------------------------
    # Title Page
    # -----------------------------

    title = doc.add_heading("AI BUSINESS DOCUMENT", level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    subtitle.add_run("Prepared By\n").bold = True
    subtitle.add_run("Autonomous AI Agent\n\n")

    subtitle.add_run("Generated On\n").bold = True
    subtitle.add_run(datetime.now().strftime("%d %B %Y %I:%M %p"))

    doc.add_page_break()

    # -----------------------------
    # Metadata
    # -----------------------------

    doc.add_heading("Document Information", level=1)

    if user_request:
        doc.add_paragraph(f"Request : {user_request}")

    doc.add_paragraph(
        f"Generated On : {datetime.now().strftime('%d %B %Y %I:%M %p')}"
    )

    if tasks:

        doc.add_heading("Execution Summary", level=2)

        for task in tasks:
            doc.add_paragraph(
                f"✓ {task.name} ({task.status})",
                style="List Bullet"
            )

    doc.add_page_break()

    # -----------------------------
    # Generated Content
    # -----------------------------

    for line in document_text.split("\n"):
        line = line.strip()
        if not line:
            continue

        # Heading 1

        if line.startswith("# "):
            heading = line.replace("#", "").strip()
            h = doc.add_heading(heading, level=1)
            h.runs[0].font.size = Pt(18)

        # Heading 2

        elif line.startswith("## "):
            heading = line.replace("##", "").strip()
            h = doc.add_heading(heading, level=2)
            h.runs[0].font.size = Pt(15)

        # Bullet Points

        elif line.startswith("- "):

            doc.add_paragraph(
                line.replace("- ", ""),
                style="List Bullet"
            )

        else:

            p = doc.add_paragraph()
            p.style.font.size = Pt(11)
            p.add_run(line)

    # -----------------------------
    # Footer
    # -----------------------------

    doc.add_page_break()
    footer = doc.add_paragraph()
    footer.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    footer.add_run(
        "Generated Automatically by Autonomous AI Agent v2"
    ).italic = True

    # -----------------------------
    # Save
    # -----------------------------

    filename = (
        f"generated_docs/document_"
        f"{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    )

    doc.save(filename)

    logger.info(f"Document generated at {filename}")

    return filename